"""JobPilot resume parser v2 — paranoid about edge cases.

Key improvements over v1:
- Multi-column PDF support via PyMuPDF block sorting
- Confidence scoring per field
- Verbatim bullet preservation (not summarized)
- Aggressive metric extraction with 15+ patterns
- Header/footer content recovery
- Table-based layout handling for DOCX
- Clear error messages for every failure mode
"""

import io
import json
import os
import re
from typing import Any

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from dotenv import load_dotenv

load_dotenv()


class ParseError(Exception):
    """Raised when resume parsing fails unrecoverably."""
    pass


# ══════════════════════════════════════════════════════════
# STEP 1: Raw text extraction (multi-format, multi-column)
# ══════════════════════════════════════════════════════════

def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF with multi-column awareness."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        raise ParseError(f"Cannot open PDF: {e}")

    if doc.page_count == 0:
        raise ParseError("PDF has 0 pages")
    if doc.page_count > 10:
        raise ParseError("PDF has too many pages for a resume (max 10). Are you sure this is a resume?")

    all_text = []
    for page in doc:
        # Strategy 1: Try block-level extraction with sorting (handles multi-column)
        blocks = page.get_text("dict", sort=True)["blocks"]
        page_lines = []
        for block in blocks:
            if block["type"] == 0:  # text block
                for line in block.get("lines", []):
                    spans_text = " ".join(span["text"] for span in line.get("spans", []))
                    if spans_text.strip():
                        page_lines.append(spans_text.strip())

        # Strategy 2: If block extraction got very little, try plain text with sort
        if len(page_lines) < 5:
            plain = page.get_text("text", sort=True)
            if len(plain.strip()) > len("\n".join(page_lines)):
                page_lines = [l.strip() for l in plain.split("\n") if l.strip()]

        if page_lines:
            all_text.append("\n".join(page_lines))

    doc.close()

    if not all_text:
        raise ParseError("PDF contains no extractable text. It may be a scanned image — try a DOCX version instead.")

    return "\n\n".join(all_text)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX: paragraphs + tables + headers."""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ParseError(f"Cannot open DOCX: {e}")

    parts = []

    # Headers (often contain name + contact info — parsers frequently miss these)
    for section in doc.sections:
        header = section.header
        if header:
            for para in header.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

    # Main body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables (some resumes use tables for layout — extract row by row)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            # Deduplicate merged cells
            seen = set()
            unique_cells = []
            for c in cells:
                if c not in seen:
                    seen.add(c)
                    unique_cells.append(c)
            if unique_cells:
                parts.append(" | ".join(unique_cells))

    if not parts:
        raise ParseError("DOCX contains no extractable text")

    return "\n".join(parts)


def _clean_text(raw: str) -> str:
    """Clean extracted text: fix whitespace, encoding artifacts, common PDF garbage."""
    text = raw
    # Remove null bytes and replacement chars
    text = text.replace("\x00", "").replace("\ufffd", "")
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces (preserve newlines)
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Fix common ligature issues
    text = text.replace("ﬁ", "fi").replace("ﬂ", "fl").replace("ﬀ", "ff")
    return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to correct extractor based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        raw = _extract_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        raw = _extract_docx(file_bytes)
    else:
        raise ParseError(
            f"Unsupported file type: .{ext}. "
            f"JobPilot accepts PDF and DOCX files only. "
            f"If you have a different format, please re-export your resume as PDF."
        )

    cleaned = _clean_text(raw)

    # Sanity check: too short = probably failed extraction
    if len(cleaned) < 100:
        raise ParseError(
            f"Extracted only {len(cleaned)} characters — this is too short for a resume. "
            f"The file may be image-based or corrupted. Try re-saving as a text-based PDF."
        )

    return cleaned


# ══════════════════════════════════════════════════════════
# STEP 2: Structure with Claude API
# ══════════════════════════════════════════════════════════

SYSTEM_PROMPT = """You are an expert resume parser. Extract structured data from the resume text below.

Return ONLY valid JSON — no markdown fences, no prose, no explanation.

CRITICAL RULES:
1. For work_history key_bullets: preserve the EXACT wording from the resume. Do NOT summarize or rephrase. The verbatim text matters for downstream cover letter generation.
2. For strongest_metrics: capture EVERY quantified achievement — numbers, percentages, rankings, scales, counts. Be aggressive. Missing a metric is worse than including a borderline one.
3. For skills: include ALL technical keywords, tools, languages, frameworks, and methodologies mentioned anywhere in the resume, even if only mentioned once.

The JSON must have exactly these fields:
{
  "name": "string or null",
  "email": "string or null",
  "phone": "string or null",
  "degree_level": "Bachelor" | "Master" | "PhD" | "Other",
  "degree_fields": ["list of degree field names"],
  "universities": ["list of university names"],
  "graduation_year": integer or null (most recent or expected),
  "gpa": float or null,
  "years_experience": float (estimate from work history dates, round to 1 decimal),
  "skills": ["ALL technical keywords, tools, languages, frameworks mentioned"],
  "companies": ["all company names worked at, in chronological order"],
  "industries": ["fintech", "e-commerce", "healthcare", "consulting", etc.],
  "languages_spoken": ["English", "Mandarin Chinese", etc.],
  "notable_achievements": ["awards, rankings, certifications — one string each"],
  "work_history": [
    {
      "company": "string",
      "title": "string",
      "duration": "string (e.g. Jul 2024 – Sep 2024)",
      "key_bullets": ["VERBATIM bullet points from the resume — do NOT rephrase"]
    }
  ],
  "strongest_metrics": ["every quantified number/percentage/ranking from the resume"],
  "extraction_confidence": {
    "name": "high" | "medium" | "low",
    "contact": "high" | "medium" | "low",
    "education": "high" | "medium" | "low",
    "experience": "high" | "medium" | "low",
    "skills": "high" | "medium" | "low",
    "metrics": "high" | "medium" | "low",
    "overall": "high" | "medium" | "low"
  }
}

Rules:
- If a field cannot be determined, use null (not empty string)
- degree_level: use the HIGHEST degree mentioned
- years_experience: sum months from all work entries
- extraction_confidence.overall: "high" if all major sections found clearly,
  "medium" if some sections ambiguous, "low" if layout was confusing or data sparse
- Do NOT invent information not present in the resume"""


def _call_claude(raw_text: str) -> dict[str, Any]:
    """Call Claude API for structured extraction."""
    api_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not api_key or api_key in ("REPLACE_ME", "sk-ant-..."):
        return _regex_fallback(raw_text)

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            system=SYSTEM_PROMPT,
            messages=[
                {"role": "user", "content": f"Parse this resume:\n\n{raw_text[:8000]}"}
            ],
        )
    except Exception as e:
        err = str(e).lower()
        if any(kw in err for kw in ["credit", "authentication", "401", "400", "api key", "rate"]):
            print(f"[parser] API unavailable ({type(e).__name__}), using regex fallback")
            return _regex_fallback(raw_text)
        raise ParseError(f"Claude API error: {e}")

    response_text = message.content[0].text.strip()

    # Strip markdown fences if present
    if response_text.startswith("```"):
        response_text = re.sub(r"^```(?:json)?\s*", "", response_text)
        response_text = re.sub(r"\s*```$", "", response_text)

    try:
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        raise ParseError(f"Claude returned invalid JSON: {e}\nFirst 300 chars: {response_text[:300]}")


# ══════════════════════════════════════════════════════════
# STEP 2b: Regex fallback (when API unavailable)
# ══════════════════════════════════════════════════════════

def _regex_fallback(raw_text: str) -> dict[str, Any]:
    """Best-effort extraction without API. Surprisingly effective for clean resumes."""
    # Email
    emails = re.findall(r'[\w.+-]+@[\w-]+\.[\w.]+', raw_text)
    email = emails[0] if emails else None

    # Phone
    phones = re.findall(r'[\d\-()+ ]{10,15}', raw_text)
    phone = phones[0].strip() if phones else None

    # Name: first non-empty line that's not an email/phone/address
    lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
    name = None
    for line in lines[:5]:
        if '@' not in line and not re.match(r'^[\d\-()+ ]+$', line) and len(line) < 60:
            name = line
            break

    # Skills: match against comprehensive keyword list
    tech_kw = [
        "Python", "SQL", "Java", "JavaScript", "TypeScript", "R", "C++", "C#", "Go", "Rust", "Scala",
        "Machine Learning", "Deep Learning", "NLP", "Computer Vision",
        "TensorFlow", "PyTorch", "scikit-learn", "XGBoost", "LightGBM",
        "pandas", "NumPy", "SciPy", "statsmodels", "Matplotlib",
        "Docker", "AWS", "GCP", "Azure", "Kubernetes",
        "Flask", "FastAPI", "Django", "React", "Next.js", "Node.js",
        "PostgreSQL", "MySQL", "MongoDB", "Redis", "Supabase",
        "Git", "Linux", "Tableau", "Power BI", "Excel",
        "A/B Testing", "ETL", "Spark", "Hadoop", "Airflow",
        "BERT", "GPT", "Sentence-BERT", "Hugging Face",
        "Random Forest", "SVM", "PCA", "DBSCAN", "ARIMA", "RNN", "CNN", "LSTM",
        "Logistic Regression", "Decision Trees", "Neural Networks", "Gradient Boosting",
    ]
    skills = [k for k in tech_kw if k.lower() in raw_text.lower()]

    # Metrics: aggressive multi-pattern extraction
    metrics = _extract_metrics_regex(raw_text)

    # Universities
    uni_patterns = [
        r"((?:University|Institute|College|School)\s+of\s+[\w\s,]+)",
        r"(Yale|Harvard|MIT|Stanford|Princeton|Columbia|UW|UCLA|Berkeley|CMU|Cornell)",
    ]
    universities = []
    for pat in uni_patterns:
        universities.extend(re.findall(pat, raw_text))
    universities = list(set(u.strip() for u in universities))

    # Degree
    degree = "Master" if any(d in raw_text for d in ["M.S.", "Master", "M.A.", "MBA"]) else \
             "PhD" if any(d in raw_text for d in ["Ph.D", "PhD", "Doctoral"]) else "Bachelor"

    # GPA
    gpa_match = re.search(r'(\d\.\d{1,2})\s*/\s*4\.0', raw_text)
    gpa = float(gpa_match.group(1)) if gpa_match else None

    # Languages
    lang_signals = {
        "English": ["english", "bilingual"],
        "Mandarin Chinese": ["mandarin", "chinese", "中文", "普通话"],
        "Spanish": ["spanish", "español"],
    }
    languages = [lang for lang, signals in lang_signals.items()
                 if any(s in raw_text.lower() for s in signals)]

    return {
        "name": name, "email": email, "phone": phone,
        "degree_level": degree,
        "degree_fields": [],
        "universities": universities,
        "graduation_year": None,
        "gpa": gpa,
        "years_experience": 0.0,
        "skills": skills,
        "companies": [],
        "industries": [],
        "languages_spoken": languages if languages else ["English"],
        "notable_achievements": [],
        "work_history": [],
        "strongest_metrics": metrics,
        "extraction_confidence": {
            "name": "high" if name else "low",
            "contact": "high" if email else "low",
            "education": "medium" if universities else "low",
            "experience": "low",  # regex can't reliably extract work history
            "skills": "high" if len(skills) >= 5 else "medium",
            "metrics": "high" if len(metrics) >= 3 else "medium" if metrics else "low",
            "overall": "medium",
        },
        "_mock": True,
    }


def _extract_metrics_regex(text: str) -> list[str]:
    """Aggressively extract every quantified achievement from text."""
    metrics = set()

    patterns = [
        # Percentages
        (r'(\d+(?:\.\d+)?)\s*%', lambda m: f"{m.group(1)}%"),
        # GPA
        (r'(\d\.\d{1,2})\s*/\s*4\.0\s*GPA', lambda m: f"{m.group(1)} GPA"),
        (r'GPA[:\s]*(\d\.\d{1,2})', lambda m: f"{m.group(1)} GPA"),
        # Rankings
        (r'top\s+(\d+)%', lambda m: f"top {m.group(1)}%"),
        (r'top\s+(\d+)\s+out\s+of', lambda m: m.group(0)),
        # Large numbers with context
        (r'(\d+(?:,\d{3})+)\s+(entrants|projects|countries|users|employees|students|interactions|comments|practitioners|sessions)',
         lambda m: f"{m.group(1)} {m.group(2)}"),
        # K+ numbers
        (r'(\d+)K\+?\s*(users|followers|interactions|comments)', lambda m: f"{m.group(1)}K+ {m.group(2)}"),
        # GB/day
        (r'(\d+[\d–-]*\d*)\s*GB(?:/day)?', lambda m: f"{m.group(1)} GB/day"),
        # AUPRC/AUC
        (r'(?:AUPRC|AUC)\s*(?:of\s*)?(\d\.\d+)', lambda m: f"{m.group(0)}"),
        # Dollar amounts
        (r'\$(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:/(?:hr|hour|month|year))?', lambda m: m.group(0)),
        # Latency
        (r'<?\s*(\d+)\s*ms\s*(?:latency|inference)', lambda m: f"<{m.group(1)}ms latency"),
        # N= sample sizes
        (r'n\s*=\s*([\d,]+\+?)', lambda m: f"n={m.group(1)}"),
        # X+ items
        (r'(\d+)\+\s+(user segments|persona|platforms|confounders|job families)',
         lambda m: f"{m.group(1)}+ {m.group(2)}"),
        # Dean's List quarters
        (r'(\d+)\s+(?:consecutive\s+)?quarters', lambda m: f"{m.group(1)} quarters"),
        # OR/CI from epidemiology
        (r'OR\s*(?:=\s*)?(\d\.\d+)', lambda m: f"OR {m.group(1)}"),
    ]

    for pattern, formatter in patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            metrics.add(formatter(match))

    return sorted(metrics)


# ══════════════════════════════════════════════════════════
# STEP 3: Validation + Confidence scoring
# ══════════════════════════════════════════════════════════

def validate_profile(profile: dict[str, Any], raw_text: str) -> dict[str, Any]:
    """Validate extracted profile. Add confidence if missing. Flag issues."""
    warnings = []

    # Critical: name must exist
    if not profile.get("name"):
        raise ParseError(
            "Could not extract your name from the resume. "
            "This usually happens with image-based PDFs or unusual formatting. "
            "Please try uploading a DOCX version instead."
        )

    # Default degree_level
    if not profile.get("degree_level"):
        profile["degree_level"] = "Bachelor"
        warnings.append("degree_level not detected, defaulted to Bachelor")

    # Ensure lists
    for field in ["skills", "companies", "industries", "universities",
                  "degree_fields", "languages_spoken", "notable_achievements",
                  "work_history", "strongest_metrics"]:
        if not isinstance(profile.get(field), list):
            profile[field] = []

    # Boost metrics from regex if Claude missed some
    regex_metrics = _extract_metrics_regex(raw_text)
    existing = set(str(m).lower() for m in profile.get("strongest_metrics", []))
    for rm in regex_metrics:
        if rm.lower() not in existing:
            profile["strongest_metrics"].append(rm)

    # Warn on low metric count
    if len(profile.get("strongest_metrics", [])) < 2:
        warnings.append(
            f"Only {len(profile.get('strongest_metrics', []))} quantified metrics found. "
            f"Consider adding numbers to your resume for stronger cover letters."
        )

    # Warn on empty work_history
    if not profile.get("work_history"):
        warnings.append("No work experience entries extracted. Check if your resume format is standard.")

    # Warn on few skills
    if len(profile.get("skills", [])) < 3:
        warnings.append("Very few technical skills detected. You may want to add a skills section to your resume.")

    # Ensure numeric fields
    for field, default in [("gpa", None), ("years_experience", 0.0), ("graduation_year", None)]:
        val = profile.get(field)
        if val is not None:
            try:
                profile[field] = float(val) if field != "graduation_year" else int(val)
            except (ValueError, TypeError):
                profile[field] = default

    # Build/supplement confidence scores
    if "extraction_confidence" not in profile or not isinstance(profile["extraction_confidence"], dict):
        profile["extraction_confidence"] = {}

    conf = profile["extraction_confidence"]

    # Auto-calculate confidence if not provided by Claude
    if "overall" not in conf:
        scores = {
            "name": "high" if profile.get("name") else "low",
            "contact": "high" if profile.get("email") else ("medium" if profile.get("phone") else "low"),
            "education": "high" if profile.get("universities") else "low",
            "experience": "high" if len(profile.get("work_history", [])) >= 2 else ("medium" if profile.get("work_history") else "low"),
            "skills": "high" if len(profile.get("skills", [])) >= 8 else ("medium" if len(profile.get("skills", [])) >= 3 else "low"),
            "metrics": "high" if len(profile.get("strongest_metrics", [])) >= 5 else ("medium" if len(profile.get("strongest_metrics", [])) >= 2 else "low"),
        }
        conf.update(scores)

        # Overall: lowest of all sections
        vals = {"high": 3, "medium": 2, "low": 1}
        avg = sum(vals.get(v, 1) for v in scores.values()) / len(scores)
        conf["overall"] = "high" if avg >= 2.5 else ("medium" if avg >= 1.8 else "low")

    # Compute numeric confidence score (0-100)
    conf_map = {"high": 100, "medium": 60, "low": 20}
    section_weights = {"name": 15, "contact": 10, "education": 15, "experience": 25, "skills": 15, "metrics": 20}
    total_score = sum(
        conf_map.get(conf.get(section, "low"), 20) * weight / 100
        for section, weight in section_weights.items()
    )
    conf["score"] = round(total_score)
    conf["needs_review"] = conf["score"] < 70 or conf.get("overall") == "low"

    profile["_warnings"] = warnings
    return profile


# ══════════════════════════════════════════════════════════
# Main entry point
# ══════════════════════════════════════════════════════════

def parse_resume(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """
    Parse a resume file into a structured profile.

    Returns structured profile with confidence scores.
    Raises ParseError with helpful user-facing messages.
    """
    # Step 1: Extract raw text
    raw_text = extract_text(file_bytes, filename)

    # Step 2: Structure with Claude (or regex fallback)
    profile = _call_claude(raw_text)

    # Step 3: Validate + enrich + score confidence
    profile = validate_profile(profile, raw_text)

    return profile
